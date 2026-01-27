#!/usr/bin/env python3
"""
Update remaining Word documents with corrected staffing model and ROI figures.
Updates:
- SYSTEM_CAPABILITIES_WIIFM.docx
- ACADEMIC_PAPER_FIGURES.docx
- ACADEMIC_PAPER_TEMPLATE.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pathlib import Path


def process_markdown_to_word(md_file, output_file, title):
    """Convert markdown file to Word document with formatting."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Add metadata
    doc.add_paragraph()
    doc.add_paragraph()
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run("HSCP Staff Rota Management System\n7 January 2026")
    meta_run.font.size = Pt(12)
    meta_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_page_break()
    
    # Read and process markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip YAML front matter
        if i == 0 and line.strip() == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 102, 153)
            elif level == 3:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(14)
                run.font.bold = True
            else:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(12)
                run.font.bold = True
            
            doc.add_paragraph()
        
        # Tables
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            
            # Skip separator line
            if len(table_lines) > 1 and '---' in table_lines[1]:
                header = table_lines[0]
                separator = table_lines[1]
                data_rows = table_lines[2:]
            else:
                header = table_lines[0]
                data_rows = table_lines[1:]
            
            # Parse header
            headers = [h.strip() for h in header.split('|')[1:-1]]
            
            # Create table
            table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Add headers
            for j, header_text in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = header_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for row_idx, row_line in enumerate(data_rows):
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                for col_idx, cell_text in enumerate(cells):
                    table.rows[row_idx + 1].cells[col_idx].text = cell_text
            
            doc.add_paragraph()
        
        # Bullet lists
        elif line.strip().startswith(('- ', '* ')):
            text = line.strip()[2:]
            add_formatted_text(doc.add_paragraph(style='List Bullet'), text)
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            add_formatted_text(doc.add_paragraph(style='List Number'), text)
        
        # Code blocks
        elif line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            para = doc.add_paragraph()
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Horizontal rules
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph('_' * 80)
        
        # Empty lines
        elif not line.strip():
            if i > 0 and i < len(lines) - 1:  # Don't add extra space at start/end
                doc.add_paragraph()
        
        # Regular paragraphs
        else:
            if line.strip():
                add_formatted_text(doc.add_paragraph(), line.strip())
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✅ Updated: {output_file}")


def add_formatted_text(paragraph, text):
    """Add text to paragraph with bold and italic formatting."""
    # Handle bold and italic markdown
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)


def main():
    """Update all remaining Word documents."""
    print("\nUpdating Word Documents with Corrected Staffing Model...")
    print("=" * 60)
    
    base_dir = Path(__file__).parent
    
    documents = [
        {
            'md': base_dir / 'SYSTEM_CAPABILITIES_WIIFM.md',
            'docx': base_dir / 'SYSTEM_CAPABILITIES_WIIFM.docx',
            'title': 'System Capabilities - What\'s In It For Me?'
        },
        {
            'md': base_dir / 'ACADEMIC_PAPER_FIGURES.md',
            'docx': base_dir / 'ACADEMIC_PAPER_FIGURES.docx',
            'title': 'Academic Paper - Figures and Tables'
        },
        {
            'md': base_dir / 'ACADEMIC_PAPER_TEMPLATE.md',
            'docx': base_dir / 'ACADEMIC_PAPER_TEMPLATE.docx',
            'title': 'Academic Paper - Staff Rota Management System'
        }
    ]
    
    for doc_info in documents:
        if doc_info['md'].exists():
            print(f"\nProcessing {doc_info['md'].name}...")
            process_markdown_to_word(
                doc_info['md'],
                doc_info['docx'],
                doc_info['title']
            )
        else:
            print(f"⚠️  Skipped: {doc_info['md'].name} (not found)")
    
    print("\n" + "=" * 60)
    print("✅ All Word documents updated successfully!")
    print("\nUpdated documents:")
    print("1. SYSTEM_CAPABILITIES_WIIFM.docx")
    print("2. ACADEMIC_PAPER_FIGURES.docx")
    print("3. ACADEMIC_PAPER_TEMPLATE.docx")
    print("\nPreviously created:")
    print("4. HSCP_12_Week_Implementation_Plan.docx")
    print("5. HSCP_ROI_Analysis.docx")
    print("6. HSCP_Business_Case_Five_Case_Model.docx")


if __name__ == '__main__':
    main()
