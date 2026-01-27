#!/usr/bin/env python3
"""
Create Word documents for HSCP Implementation Plan and ROI Analysis
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def process_markdown_to_word(markdown_file, output_file, title):
    """Convert markdown to Word document"""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title page
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Health and Social Care Partnership\\n5 Care Homes, 200 Staff')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 84, 106)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('7 January 2026')
    run.font.size = Pt(12)
    
    add_page_break(doc)
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process content
    lines = content.split('\\n')
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []
    skip_next = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if skip_next:
            skip_next = False
            i += 1
            continue
        
        # Skip YAML frontmatter
        if line.strip() == '---':
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code:
                # End code block
                if code_lines:
                    code_para = doc.add_paragraph('\\n'.join(code_lines))
                    code_para.style = 'Intense Quote'
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip():
            if not in_table:
                in_table = True
                table_lines = []
            
            # Check if separator line
            if set(line.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                # Skip separator
                i += 1
                continue
            
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_lines:
                create_word_table(doc, table_lines)
            table_lines = []
            in_table = False
            # Don't increment i, process this line normally
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            if heading_text and level <= 3:
                h = doc.add_heading(heading_text, level)
                # Add some spacing
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
        # Handle lists
        elif line.strip().startswith(('- ', '* ', '+ ')):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text)
        elif re.match(r'^\\d+\\.', line.strip()):
            text = re.sub(r'^\\d+\\.\\s*', '', line.strip())
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, text)
        # Regular paragraphs
        elif line.strip():
            para = doc.add_paragraph()
            add_formatted_text(para, line)
        else:
            # Empty line for spacing
            if i > 0 and i < len(lines) - 1:
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f'✅ Created: {output_file}')

def create_word_table(doc, lines):
    """Create a table from markdown table lines"""
    if not lines:
        return
    
    # Parse table
    rows = []
    for line in lines:
        if not line.strip():
            continue
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first and last cells
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.rows[i].cells[j]
            if j < len(row):
                cell.text = row[j]
                # Make header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)

def add_formatted_text(para, text):
    """Add formatted text to paragraph"""
    # Handle bold (**text**)
    parts = re.split(r'(\\*\\*.*?\\*\\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Handle inline code and other formatting
            text_clean = part.replace('`', '')
            if text_clean:
                para.add_run(text_clean)

if __name__ == '__main__':
    print('Creating HSCP Word Documents...\\n')
    
    # Create Implementation Plan
    print('Processing 12-Week Implementation Plan...')
    process_markdown_to_word(
        'HSCP_12_WEEK_IMPLEMENTATION_PLAN.md',
        'HSCP_12_Week_Implementation_Plan.docx',
        'HSCP 12-Week Implementation Plan'
    )
    
    # Create ROI Analysis
    print('\\nProcessing ROI Analysis...')
    process_markdown_to_word(
        'HSCP_REVISED_ROI_ANALYSIS.md',
        'HSCP_ROI_Analysis.docx',
        'Staff Rota System - ROI Analysis'
    )
    
    print('\\n✅ All Word documents created successfully!')
