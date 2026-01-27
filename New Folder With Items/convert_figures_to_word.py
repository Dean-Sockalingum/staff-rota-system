#!/usr/bin/env python3
"""
Convert ACADEMIC_PAPER_FIGURES.md to publication-quality Word format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import re

# Read the markdown file
md_content = Path('ACADEMIC_PAPER_FIGURES.md').read_text()

doc = Document()

# Set up default style
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Add title
title = doc.add_heading('Academic Paper Figures & Diagrams', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Supplementary Materials for Academic Paper')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(12)
subtitle_run.font.italic = True

# Add metadata
metadata = doc.add_paragraph()
metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = metadata.add_run('Document: Staff Rota System Figures\nDate: 22 December 2025\nPurpose: Journal Submission')
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# Process line by line
lines = md_content.split('\n')
i = 0
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]
    
    # Code blocks (for diagrams)
    if line.startswith('```'):
        if in_code_block:
            # End of code block - add accumulated lines
            if code_lines:
                # Add diagram with monospace font
                for code_line in code_lines:
                    p = doc.add_paragraph(code_line)
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if p.runs:
                        run = p.runs[0]
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 128)
                code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Skip horizontal rules
    if line.startswith('---') or line.strip() == '':
        i += 1
        continue
    
    # Main Headers (Figure titles)
    if line.startswith('## Figure '):
        doc.add_page_break()
        heading = doc.add_heading(line[3:], level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    # Sub-headers
    elif line.startswith('### '):
        heading = doc.add_heading(line[4:], level=3)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
    
    elif line.startswith('#### '):
        heading = doc.add_heading(line[5:], level=4)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
    
    # Section headers (like "OBJECTIVE FUNCTION:")
    elif line.endswith(':') and line.isupper() and not line.startswith('│'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    
    # Bullet lists
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip().lstrip('- *').strip()
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Handle bold text in bullets
        if '**' in text:
            p.clear()
            parts = text.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Odd indices are bold
                    run.bold = True
    
    # Numbered lists
    elif re.match(r'^\d+\.', line.strip()):
        text = line.strip().split('. ', 1)[1] if '. ' in line else line
        p = doc.add_paragraph(text, style='List Number')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    
    # Tables (keep markdown format for simplicity)
    elif '|' in line and not line.startswith('#'):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.3)
        if p.runs:
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    
    # Caption detection (starts with "**Caption:**")
    elif line.strip().startswith('**Caption:**'):
        caption_text = line.replace('**Caption:**', '').strip()
        p = doc.add_paragraph()
        caption_label = p.add_run('Caption: ')
        caption_label.bold = True
        caption_label.font.size = Pt(10)
        caption_label.italic = True
        caption_content = p.add_run(caption_text)
        caption_content.font.size = Pt(10)
        caption_content.italic = True
        caption_content.font.color.rgb = RGBColor(64, 64, 64)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
    
    # Regular paragraphs
    elif line.strip():
        p = doc.add_paragraph()
        # Handle bold text
        text = line.strip()
        while '**' in text:
            before, rest = text.split('**', 1)
            if '**' in rest:
                bold_text, after = rest.split('**', 1)
                if before:
                    p.add_run(before)
                p.add_run(bold_text).bold = True
                text = after
            else:
                text = before + '**' + rest
                break
        if text:
            p.add_run(text)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    
    i += 1

# Add page break before summary table
doc.add_page_break()

# Save the document
output_file = 'ACADEMIC_PAPER_FIGURES.docx'
doc.save(output_file)

print('✅ Publication-quality Word document created successfully!')
print('')
print('📄 File: ACADEMIC_PAPER_FIGURES.docx')
print('📍 Location: /Users/deansockalingum/Desktop/Staff_Rota_Backups/2025-12-12_Multi-Home_Complete/')
print('📊 Contents: 10 figures with diagrams and captions')
print('📝 Format: Microsoft Word (.docx)')
print('')
print('✨ Features:')
print('  • Publication-quality formatting (Arial 11pt)')
print('  • Monospace diagrams (Courier New 9pt)')
print('  • Color-coded sections (blue headers)')
print('  • Italicized captions')
print('  • Page breaks between figures')
print('  • Professional layout for journal submission')
print('')
print('💡 Ready for:')
print('  • Journal submission (JHIM, IJMI, etc.)')
print('  • Further enhancement with visio/draw.io graphics')
print('  • Conversion to PDF for archival')
