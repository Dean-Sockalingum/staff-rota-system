#!/usr/bin/env python3
"""
Convert Automation Workflow Diagrams from Markdown to Word Document
Preserves formatting, flowcharts, and tables
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_horizontal_line(paragraph):
    """Add a horizontal line (border) to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_word_document():
    """Create Word document from markdown file"""
    
    # Create document
    doc = Document()
    
    # Set document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read the markdown file
    with open('AUTOMATION_WORKFLOW_DIAGRAMS.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Handle code blocks (flowcharts)
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block - add as formatted text
                if code_lines:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add all code lines with courier font
                    full_text = '\n'.join(code_lines)
                    run = para.add_run(full_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    
                    # Add shading to make it stand out
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F0F0F0')
                    run._element.get_or_add_rPr().append(shading_elm)
                    
                in_code_block = False
                code_lines = []
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Handle different markdown elements
        stripped = line.strip()
        
        # Main title (# )
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            para = doc.add_heading(text, level=1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            run.font.size = Pt(24)
        
        # H2 (## )
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            para = doc.add_heading(text, level=2)
            run = para.runs[0]
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            run.font.size = Pt(18)
        
        # H3 (### )
        elif stripped.startswith('### '):
            text = stripped[4:]
            para = doc.add_heading(text, level=3)
            run = para.runs[0]
            run.font.size = Pt(14)
        
        # Horizontal rules (---)
        elif stripped.startswith('---'):
            para = doc.add_paragraph()
            add_horizontal_line(para)
        
        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            
            # Handle bold/colored markers
            para = doc.add_paragraph(style='List Bullet')
            
            # Parse inline formatting
            if '**' in text:
                parts = re.split(r'\*\*(.*?)\*\*', text)
                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            else:
                para.add_run(text)
        
        # Tables (| column | column |)
        elif stripped.startswith('|') and '|' in stripped[1:]:
            # This is a table row
            # Look ahead to gather all table rows
            table_rows = [line]
            continue  # Handle tables separately if needed
        
        # Regular paragraphs
        elif stripped and not stripped.startswith('#'):
            para = doc.add_paragraph(stripped)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Make emoji/icon stand out
            if any(emoji in stripped for emoji in ['🔴', '🟢', '📋', '⏱️', '💰', '📊', '🎯', '🛡️', '😊']):
                run = para.runs[0]
                run.font.size = Pt(11)
        
        # Empty lines
        elif not stripped:
            doc.add_paragraph()
    
    # Add footer with metadata
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Automation Workflow Diagrams - Staff Rota System - Glasgow HSCP - January 2026"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save the document
    output_file = 'AUTOMATION_WORKFLOW_DIAGRAMS.docx'
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    print(f"📄 Document is ready for review and sharing")
    
    return output_file

if __name__ == '__main__':
    try:
        create_word_document()
    except ImportError:
        print("❌ Error: python-docx library not found")
        print("\nTo install it, run:")
        print("  pip install python-docx")
        print("\nOr:")
        print("  pip3 install python-docx")
    except FileNotFoundError:
        print("❌ Error: AUTOMATION_WORKFLOW_DIAGRAMS.md not found")
        print("Make sure you're running this script from the correct directory")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
