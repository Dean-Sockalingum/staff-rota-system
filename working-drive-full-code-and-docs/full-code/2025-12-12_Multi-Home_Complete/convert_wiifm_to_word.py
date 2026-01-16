#!/usr/bin/env python3
"""
Convert SYSTEM_CAPABILITIES_WIIFM.md to a professionally formatted Word document.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_styled_document():
    """Create a Word document with professional styling."""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(51, 153, 255)  # Light blue
    
    return doc

def add_markdown_content(doc, md_file_path):
    """Parse markdown and add formatted content to Word document."""
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines unless in a table
        if not line and not in_table:
            i += 1
            continue
        
        # Title (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=0)
            p.style = 'Title'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading 1 (## )
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            doc.add_heading(text, level=1)
            doc.add_paragraph()  # Add spacing
        
        # Heading 2 (### )
        elif line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            doc.add_heading(text, level=2)
        
        # Heading 3 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=3)
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph()
        
        # Bullet list
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            text = line[2:].strip()
            # Handle bold/italic markdown
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold markers for now
            text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic markers
            text = re.sub(r'`(.+?)`', r'\1', text)  # Remove code markers
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered list
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Table detection
        elif '|' in line and not in_table:
            # Start collecting table data
            in_table = True
            table_data = [line]
        
        elif in_table:
            if '|' in line:
                table_data.append(line)
            else:
                # End of table, process it
                add_table_to_doc(doc, table_data)
                table_data = []
                in_table = False
                continue  # Process this line again as non-table content
        
        # Quote/callout (> )
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            p.style = 'Intense Quote'
        
        # Regular paragraph
        else:
            # Skip table separator lines
            if re.match(r'^[\s\|:-]+$', line):
                i += 1
                continue
            
            text = line.strip()
            if text:
                # Handle inline formatting
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                text = re.sub(r'`(.+?)`', r'\1', text)
                
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Process any remaining table
    if in_table and table_data:
        add_table_to_doc(doc, table_data)

def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document."""
    if len(table_data) < 2:
        return
    
    # Parse table rows
    rows = []
    for line in table_data:
        if '|' in line and not re.match(r'^[\s\|:-]+$', line):
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first/last cells from markdown table format
            cells = [c for c in cells if c]
            if cells:
                rows.append(cells)
    
    if len(rows) < 2:
        return
    
    # Create table
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            
            # Make header row bold
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                cell._element.get_or_add_tcPr().append(
                    cell._element._new_tblPr()
                )
                # Set header background color
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '0066CC')  # Blue background
                cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()  # Add spacing after table

def main():
    """Main conversion function."""
    print("Converting SYSTEM_CAPABILITIES_WIIFM.md to Word document...")
    
    # Create document
    doc = create_styled_document()
    
    # Add content
    md_file = 'SYSTEM_CAPABILITIES_WIIFM.md'
    add_markdown_content(doc, md_file)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Save document
    output_file = 'SYSTEM_CAPABILITIES_WIIFM.docx'
    doc.save(output_file)
    
    print(f"✅ Successfully created {output_file}")
    print(f"📄 Document is ready for distribution to stakeholders")

if __name__ == '__main__':
    main()
